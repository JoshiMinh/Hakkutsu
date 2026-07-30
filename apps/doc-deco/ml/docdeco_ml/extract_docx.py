from __future__ import annotations

import hashlib
import re
from collections.abc import Iterator
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .records import ParagraphRecord
from .weak_labels import infer_zone, weak_label


def _iter_blocks(parent: DocumentObject | _Cell) -> Iterator[tuple[Paragraph, bool]]:
    element = parent.element.body if isinstance(parent, DocumentObject) else parent._tc
    for child in element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent), isinstance(parent, _Cell)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            seen_cells: set[int] = set()
            for row in table.rows:
                for cell in row.cells:
                    marker = id(cell._tc)
                    if marker in seen_cells:
                        continue
                    seen_cells.add(marker)
                    yield from _iter_blocks(cell)


def _length_points(value) -> float:
    return float(value.pt) if value is not None else 0.0


def _effective_font_size(paragraph: Paragraph, run) -> float:
    if run.font.size:
        return run.font.size.pt
    if paragraph.style and paragraph.style.font.size:
        return paragraph.style.font.size.pt
    return 0.0


def _paragraph_features(paragraph: Paragraph, text: str, *, in_table: bool, container: str) -> dict[str, float]:
    runs = [run for run in paragraph.runs if run.text]
    total_chars = max(sum(len(run.text) for run in runs), 1)
    sizes = [_effective_font_size(paragraph, run) for run in runs]
    sizes = [size for size in sizes if size > 0]
    ppr = paragraph._p.pPr
    numpr = ppr.numPr if ppr is not None else None
    ilvl = numpr.find(qn("w:ilvl")) if numpr is not None else None
    outline = ppr.find(qn("w:outlineLvl")) if ppr is not None else None
    style_name = paragraph.style.name if paragraph.style else ""
    letters = [char for char in text if char.isalpha()]
    uppercase = sum(1 for char in letters if char.isupper()) / max(len(letters), 1)
    fmt = paragraph.paragraph_format
    alignment = str(paragraph.alignment or "").upper()
    return {
        "char_length": min(len(text) / 500.0, 1.0),
        "word_count": min(len(text.split()) / 100.0, 1.0),
        "font_size_mean": (sum(sizes) / len(sizes) / 30.0) if sizes else 0.0,
        "font_size_max": (max(sizes) / 30.0) if sizes else 0.0,
        "bold_ratio": sum(len(run.text) for run in runs if run.bold) / total_chars,
        "italic_ratio": sum(len(run.text) for run in runs if run.italic) / total_chars,
        "uppercase_ratio": uppercase,
        "is_centered": float("CENTER" in alignment),
        "is_justified": float("JUSTIFY" in alignment),
        "is_right_aligned": float("RIGHT" in alignment),
        "first_line_indent": min(abs(_length_points(fmt.first_line_indent)) / 72.0, 2.0),
        "left_indent": min(abs(_length_points(fmt.left_indent)) / 144.0, 2.0),
        "space_before": min(_length_points(fmt.space_before) / 72.0, 2.0),
        "space_after": min(_length_points(fmt.space_after) / 72.0, 2.0),
        "has_numbering": float(numpr is not None),
        "numbering_depth": float(int(ilvl.get(qn("w:val"))) + 1 if ilvl is not None else 0),
        "outline_level": float(int(outline.get(qn("w:val"))) if outline is not None else 9),
        "is_in_table": float(in_table),
        "is_header": float(container == "header"),
        "is_footer": float(container == "footer"),
        "page_break_before": float(bool(fmt.page_break_before)),
        "style_is_heading": float(style_name.casefold().startswith("heading")),
        "style_is_caption": float("caption" in style_name.casefold()),
        "style_is_list": float("list" in style_name.casefold()),
    }


def _document_id(path: Path) -> str:
    stat = path.stat()
    payload = f"{path.resolve()}:{stat.st_size}:{stat.st_mtime_ns}"
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:20]


def extract_docx(path: Path) -> list[ParagraphRecord]:
    document = Document(path)
    doc_id = _document_id(path)
    raw: list[tuple[Paragraph, bool, str]] = [
        (paragraph, in_table, "body") for paragraph, in_table in _iter_blocks(document)
    ]
    for section in document.sections:
        raw.extend((paragraph, False, "header") for paragraph in section.header.paragraphs)
        raw.extend((paragraph, False, "footer") for paragraph in section.footer.paragraphs)

    records: list[ParagraphRecord] = []
    zone = "cover"
    for paragraph, in_table, container in raw:
        text = re.sub(r"\s+", " ", paragraph.text.strip())
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        features = _paragraph_features(paragraph, text, in_table=in_table, container=container)
        if container == "body":
            zone = infer_zone(text, style, zone, len(records))
        role, confidence, source = weak_label(
            text=text, style=style, zone=zone, features=features, container=container
        )
        records.append(ParagraphRecord(
            document_id=doc_id,
            paragraph_id=f"p-{len(records)}",
            index=len(records),
            text=text,
            label=role.value,
            label_source=source,
            label_confidence=confidence,
            zone=zone if container == "body" else container,
            style_name=style,
            features=features,
        ))
    for index, record in enumerate(records):
        record.previous_text = records[index - 1].text if index else ""
        record.next_text = records[index + 1].text if index + 1 < len(records) else ""
        if record.zone != "cover":
            continue
        previous = records[index - 1] if index else None
        following = records[index + 1] if index + 1 < len(records) else None
        uppercase = record.features.get("uppercase_ratio", 0.0)
        if (
            previous
            and previous.label == "cover_institution"
            and following
            and following.label == "cover_faculty"
            and uppercase >= 0.7
        ):
            record.label = "cover_institution"
            record.label_source = "cover_sequence"
            record.label_confidence = 0.97
        elif (
            previous
            and previous.label == "cover_project_type"
            and following
            and following.label == "author_metadata"
            and uppercase >= 0.7
        ):
            record.label = "document_title"
            record.label_source = "cover_sequence"
            record.label_confidence = 0.97
    return records
