from __future__ import annotations

import json
import sqlite3
import tempfile
import unittest
from contextlib import closing
from pathlib import Path

from docx import Document

from docdeco_ml.build_dataset import _feedback_records, _split, build_dataset
from docdeco_ml.extract_docx import extract_docx
from docdeco_ml.labels import LABELS
from docdeco_ml.records import FEATURE_NAMES
from docdeco_ml.synthetic import generate_synthetic
from docdeco_ml.weak_labels import weak_label


class DataPipelineTests(unittest.TestCase):
    def test_synthetic_corpus_covers_every_role(self) -> None:
        records = list(generate_synthetic(100, seed=20260729))
        self.assertEqual(set(LABELS), {record.label for record in records})
        self.assertTrue(all(len(record.feature_vector()) == len(FEATURE_NAMES) for record in records))

    def test_toc_figure_entry_is_not_real_caption(self) -> None:
        role, confidence, source = weak_label(
            text="Hình 2.1. Kiến trúc tổng thể ........ 25",
            zone="list_of_figures",
            style="Normal",
            features={},
            container="body",
        )
        self.assertEqual(role.value, "list_of_figures_entry")
        self.assertGreaterEqual(confidence, 0.9)
        self.assertEqual(source, "zone_context")

    def test_extractor_keeps_text_and_layout_features(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "sample.docx"
            document = Document()
            document.add_heading("CHƯƠNG 1. TỔNG QUAN", level=1)
            document.add_paragraph("Nội dung mô tả hệ thống và phạm vi của đề tài.")
            document.add_paragraph("Hình 1.1. Kiến trúc hệ thống", style="Caption")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Thuộc tính"
            table.cell(0, 1).text = "Giá trị"
            document.save(path)

            records = extract_docx(path)
            self.assertTrue(any(record.label == "heading_1" for record in records))
            self.assertTrue(any(record.label == "figure_caption" for record in records))
            self.assertTrue(any(record.features["is_in_table"] == 1 for record in records))
            self.assertTrue(all(len(record.feature_vector()) == len(FEATURE_NAMES) for record in records))

    def test_feedback_is_imported_as_human_label(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            db_path = Path(directory) / "docdeco.db"
            with closing(sqlite3.connect(db_path)) as db:
                db.execute(
                    "CREATE TABLE feedback ("
                    "id INTEGER PRIMARY KEY, document_id TEXT, paragraph_id TEXT, "
                    "text TEXT, predicted_role TEXT, corrected_role TEXT, "
                    "context_json TEXT, created_at TEXT)"
                )
                db.execute(
                    "INSERT INTO feedback VALUES (1, 'doc', 'p1', '1.2 Mục tiêu', "
                    "'body', 'heading2', ?, '')",
                    (json.dumps({"layout_features": {"bold_ratio": 1}}),),
                )
                db.commit()
            records = _feedback_records(db_path)
            self.assertEqual(len(records), 1)
            self.assertEqual(records[0].label, "heading_2")
            self.assertEqual(records[0].label_source, "human_feedback")
            self.assertEqual(records[0].features["bold_ratio"], 1)

    def test_dataset_split_is_document_level(self) -> None:
        self.assertEqual(_split("same-document"), _split("same-document"))
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            manifest = build_dataset(root / "raw", root / "out", 30, 20260729)
            self.assertEqual(sum(manifest["documents"].values()), 30)
            split_documents: dict[str, set[str]] = {}
            for split in ("train", "validation", "test"):
                ids = set()
                for line in (root / "out" / f"{split}.jsonl").read_text(encoding="utf-8").splitlines():
                    ids.add(json.loads(line)["document_id"])
                split_documents[split] = ids
            self.assertFalse(split_documents["train"] & split_documents["validation"])
            self.assertFalse(split_documents["train"] & split_documents["test"])
            self.assertFalse(split_documents["validation"] & split_documents["test"])


if __name__ == "__main__":
    unittest.main()
